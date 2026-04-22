"""
Export service for generating professional DOCX and PDF files from proposal content.
Produces polished, government-ready documents with tables, matrices, and proper formatting.
"""

import base64
import io
import logging
import re
import urllib.parse
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from html import unescape as html_unescape

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph as RLParagraph,
    Spacer,
    Table as RLTable,
    TableStyle,
    PageBreak,
    KeepTogether,
)

logger = logging.getLogger(__name__)

# Color palette
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
NAVY_HEX = "1B2A4A"
ACCENT = RGBColor(0x10, 0xB9, 0x81)
ACCENT_HEX = "10B981"
DARK_TEXT = RGBColor(0x37, 0x41, 0x51)
GRAY_TEXT = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BG = "F0FDF4"
HEADER_BG = "1B2A4A"
SUB_HEADER_BG = "059669"
LIGHT_GRAY_BG = "F3F4F6"
BORDER_COLOR = RGBColor(0xD1, 0xD5, 0xDB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
UPLOADS_DIR = Path(__file__).resolve().parents[1] / "data" / "uploads"

# Page content width for PDF (letter 8.5" - 0.75" left - 0.75" right margins)
PDF_CONTENT_WIDTH = 7.0 * inch


def _normalize_hex_color(value: str, fallback: str) -> str:
    """Return a ReportLab-friendly hex color string."""
    if not value:
        return fallback
    value = value.strip()
    if re.match(r"^#?[0-9A-Fa-f]{6}$", value):
        return value if value.startswith("#") else f"#{value}"
    return fallback


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color="D1D5DB", size="4"):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _add_formatted_paragraph(doc, text: str, size=11, color=DARK_TEXT, bold=False,
                              italic=False, alignment=None, space_before=0, space_after=6,
                              font_name="Calibri"):
    """Add a paragraph with specific formatting."""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
            run.bold = bold

    for run in para.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = font_name
        if italic:
            run.italic = True
    return para


def _strip_leading_title(content: str, section_title: str) -> str:
    """Remove duplicate section title from the start of content."""
    if not content or not section_title:
        return content
    title_lower = section_title.lower().strip()

    heading_pattern = rf'\s*<h[1-6][^>]*>\s*{re.escape(title_lower)}\s*</h[1-6]>\s*'
    content = re.sub(heading_pattern, '', content, count=1, flags=re.IGNORECASE).strip()

    text_only = re.sub(r'<[^>]+>', '', content[:200]).strip()
    if text_only.lower().startswith(title_lower):
        inner_pattern = rf'(<p[^>]*>(?:\s*<(?:strong|b|em|i)[^>]*>)?\s*){re.escape(title_lower)}(\s*(?:</(?:strong|b|em|i)>)?\s*)'
        cleaned = re.sub(inner_pattern, r'\1', content, count=1, flags=re.IGNORECASE)
        if cleaned != content:
            cleaned = re.sub(r'<p[^>]*>\s*(?:<(?:strong|b|em|i)[^>]*>\s*</(?:strong|b|em|i)>\s*)*</p>\s*', '', cleaned)
            content = cleaned.strip()

    return content


def _add_styled_table(doc, headers: List[str], rows: List[List[str]],
                       header_bg=HEADER_BG, alt_row_bg=LIGHT_GRAY_BG):
    """Add a professionally styled table."""
    if not headers or not rows:
        return None

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        _set_cell_shading(cell, header_bg)
        _set_cell_borders(cell, NAVY_HEX)

    for row_idx, row_data in enumerate(rows):
        for col_idx in range(num_cols):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT
            run.font.name = "Calibri"
            _set_cell_borders(cell)
            if row_idx % 2 == 1:
                _set_cell_shading(cell, alt_row_bg)

    return table


def _add_html_table_to_docx(doc, headers: List[str], rows: List[Dict],
                             header_bg=HEADER_BG, alt_row_bg=LIGHT_GRAY_BG, summary_bg=SUB_HEADER_BG):
    """Render an HTML table as a professional DOCX table."""
    if not headers and not rows:
        return None

    num_cols = len(headers) if headers else (len(rows[0]['cells']) if rows else 0)
    if num_cols == 0:
        return None

    total_rows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=total_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_offset = 0
    if headers:
        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            run.font.name = "Calibri"
            _set_cell_shading(cell, header_bg)
            _set_cell_borders(cell, NAVY_HEX)
        row_offset = 1

    for row_idx, row_data in enumerate(rows):
        cells_data = row_data.get('cells', [])
        is_summary = row_data.get('is_summary', False)
        for col_idx in range(num_cols):
            cell = table.rows[row_offset + row_idx].cells[col_idx]
            cell_text = cells_data[col_idx] if col_idx < len(cells_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            if cell_text.startswith('$') or (col_idx >= num_cols - 2 and re.match(r'^[\d,.$]+$', cell_text.strip())):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            _set_cell_borders(cell)

            if is_summary:
                run.bold = True
                run.font.color.rgb = NAVY
                _set_cell_shading(cell, "E8F5E9")
            else:
                run.font.color.rgb = DARK_TEXT
                if row_idx % 2 == 1:
                    _set_cell_shading(cell, alt_row_bg)

    return table


def _detect_table_content(lines: List[str]) -> Optional[Tuple[List[str], List[List[str]]]]:
    """Detect if a block of lines contains table-like data."""
    pipe_lines = [l for l in lines if '|' in l and l.count('|') >= 2]
    if len(pipe_lines) >= 2:
        headers = []
        rows = []
        for i, line in enumerate(pipe_lines):
            cells = [c.strip() for c in line.strip('|').split('|')]
            cells = [c for c in cells if c and not re.match(r'^[-:]+$', c)]
            if not cells:
                continue
            if not headers:
                headers = cells
            else:
                rows.append(cells)
        if headers and rows:
            return headers, rows

    kv_lines = []
    for line in lines:
        if ':' in line and not line.startswith('#'):
            parts = line.split(':', 1)
            if len(parts) == 2 and len(parts[0].strip()) > 1 and len(parts[1].strip()) > 1:
                kv_lines.append((parts[0].strip(), parts[1].strip()))
    if len(kv_lines) >= 3:
        return ["Item", "Details"], [[k, v] for k, v in kv_lines]

    return None


def _detect_risk_register(text: str) -> Optional[Tuple[List[str], List[List[str]]]]:
    """Detect risk register format."""
    pattern = r'(?:Risk|Issue)\s*:\s*(.+?)\s*\|\s*(?:Probability|Likelihood)\s*:\s*(\w+)\s*\|\s*Impact\s*:\s*(\w+)\s*\|\s*Mitigation\s*:\s*(.+?)(?:\s*\|\s*Contingency\s*:\s*(.+?))?(?:\n|$)'
    matches = re.findall(pattern, text, re.IGNORECASE)
    if matches:
        headers = ["Risk", "Probability", "Impact", "Mitigation"]
        if any(m[4] for m in matches):
            headers.append("Contingency")
        rows = []
        for m in matches:
            row = [m[0].strip(), m[1].strip(), m[2].strip(), m[3].strip()]
            if len(headers) == 5:
                row.append(m[4].strip() if m[4] else "")
            rows.append(row)
        return headers, rows
    return None


def generate_docx(proposal: Dict) -> io.BytesIO:
    """Generate a professional Word document from proposal sections."""
    try:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

        proposal_title = proposal.get("proposal_title", "Government Proposal")
        vendor_name = proposal.get("vendor_name", "")
        sections = proposal.get("sections", {})

        meta = proposal.get("metadata", {})
        proposal_type = meta.get("proposal_type", "Government Contract Proposal")
        agency = meta.get("agency", "")
        contracting_office = meta.get("contracting_office", "")
        solicitation_number = meta.get("solicitation_number", "")
        submission_date = meta.get("submission_date", "")
        cage_code = meta.get("cage_code", "")
        duns_number = meta.get("duns_number", "")
        naics_codes = meta.get("naics_codes", [])
        poc_name = meta.get("poc_name", "")
        poc_title_str = meta.get("poc_title", "")
        poc_email = meta.get("poc_email", "")
        poc_phone = meta.get("poc_phone", "")
        display_date = submission_date if submission_date else datetime.now().strftime('%B %d, %Y')

        company_logo = proposal.get("company_logo", "")
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()

        footer_table = footer.add_table(rows=1, cols=3, width=Inches(7))
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        left_cell = footer_table.rows[0].cells[0]
        left_cell.text = ""
        left_cell.width = Inches(2.3)
        lp = left_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        if company_logo and company_logo.startswith("data:image"):
            try:
                _, b64data = company_logo.split(",", 1)
                logo_bytes = base64.b64decode(b64data)
                logo_stream = io.BytesIO(logo_bytes)
                from PIL import Image as PILImage
                _pil_img = PILImage.open(io.BytesIO(logo_bytes))
                _iw, _ih = _pil_img.size
                _max_w, _max_h = 0.8, 0.35
                _scale = min(_max_w / (_iw / 96.0), _max_h / (_ih / 96.0), 1.0)
                _fit_w = (_iw / 96.0) * _scale
                _fit_h = (_ih / 96.0) * _scale
                lr = lp.add_run()
                lr.add_picture(logo_stream, width=Inches(_fit_w), height=Inches(_fit_h))
            except Exception:
                lr = lp.add_run(vendor_name or "")
                lr.font.size = Pt(7)
                lr.font.color.rgb = GRAY_TEXT
                lr.font.name = "Calibri"
        else:
            lr = lp.add_run(vendor_name or "")
            lr.font.size = Pt(7)
            lr.font.color.rgb = GRAY_TEXT
            lr.font.name = "Calibri"

        center_cell = footer_table.rows[0].cells[1]
        center_cell.text = ""
        center_cell.width = Inches(2.4)
        cp = center_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cr = cp.add_run("Page ")
        cr.font.size = Pt(8)
        cr.font.color.rgb = GRAY_TEXT
        cr.font.name = "Calibri"
        cr2_elem = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="16"/><w:color w:val="6B7280"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>')
        instrText = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="16"/><w:color w:val="6B7280"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
        fldChar2 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="16"/><w:color w:val="6B7280"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>')
        cp._p.append(cr2_elem)
        cp._p.append(instrText)
        cp._p.append(fldChar2)

        right_cell = footer_table.rows[0].cells[2]
        right_cell.text = ""
        right_cell.width = Inches(2.3)
        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(0)
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(f"Confidential  |  {vendor_name or ''}")
        rr.font.size = Pt(7)
        rr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        rr.font.bold = False
        rr.font.name = "Calibri"

        for row in footer_table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="none" w:sz="0" w:space="0"/>'
                    f'  <w:left w:val="none" w:sz="0" w:space="0"/>'
                    f'  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
                    f'  <w:right w:val="none" w:sz="0" w:space="0"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(tcBorders)

        # --- Cover Page ---
        banner = doc.add_table(rows=1, cols=1)
        banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        bc = banner.rows[0].cells[0]
        bc.text = ""
        _set_cell_shading(bc, NAVY_HEX)
        bc.width = Inches(7)
        bp = bc.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(24)
        bp.paragraph_format.space_after = Pt(24)
        tr = bp.add_run(proposal_title.upper())
        tr.bold = True
        tr.font.size = Pt(24)
        tr.font.color.rgb = WHITE
        tr.font.name = "Calibri"

        accent_bar = doc.add_table(rows=1, cols=1)
        accent_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
        ac = accent_bar.rows[0].cells[0]
        ac.text = ""
        _set_cell_shading(ac, ACCENT_HEX)
        ac.width = Inches(7)
        ap = ac.paragraphs[0]
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ap.paragraph_format.space_before = Pt(6)
        ap.paragraph_format.space_after = Pt(6)
        ar = ap.add_run(proposal_type)
        ar.font.size = Pt(13)
        ar.font.color.rgb = WHITE
        ar.font.name = "Calibri"
        ar.bold = True

        doc.add_paragraph("")

        company_logo = proposal.get("company_logo", "")
        if company_logo and company_logo.startswith("data:image"):
            try:
                header, b64data = company_logo.split(",", 1)
                logo_bytes = base64.b64decode(b64data)
                logo_stream = io.BytesIO(logo_bytes)
                from PIL import Image as PILImage
                _pil_img = PILImage.open(io.BytesIO(logo_bytes))
                _iw, _ih = _pil_img.size
                _max_w, _max_h = 2.0, 1.0
                _scale = min(_max_w / (_iw / 96.0), _max_h / (_ih / 96.0), 1.0)
                _fit_w = (_iw / 96.0) * _scale
                _fit_h = (_ih / 96.0) * _scale
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_para.paragraph_format.space_after = Pt(8)
                run = logo_para.add_run()
                run.add_picture(logo_stream, width=Inches(_fit_w), height=Inches(_fit_h))
            except Exception as logo_err:
                logger.warning("Failed to embed logo in DOCX: %s", logo_err)
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lr = logo_para.add_run("[Company Logo]")
                lr.font.size = Pt(10)
                lr.font.color.rgb = GRAY_TEXT
                lr.italic = True
        else:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.paragraph_format.space_after = Pt(4)
            lr = logo_para.add_run("[Upload logo in Business Profile]")
            lr.font.size = Pt(10)
            lr.font.color.rgb = GRAY_TEXT
            lr.italic = True
            lr.font.name = "Calibri"

        doc.add_paragraph("")

        info_table = doc.add_table(rows=0, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.columns[0].width = Inches(2.2)
        info_table.columns[1].width = Inches(4.3)

        def _add_info_row(label, value, value_bold=False, value_size=11):
            row = info_table.add_row()
            lc = row.cells[0]
            lc.text = ""
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lp.paragraph_format.space_before = Pt(4)
            lp.paragraph_format.space_after = Pt(4)
            lrun = lp.add_run(label)
            lrun.font.size = Pt(10)
            lrun.font.color.rgb = GRAY_TEXT
            lrun.font.name = "Calibri"
            lrun.bold = True
            vc = row.cells[1]
            vc.text = ""
            vp = vc.paragraphs[0]
            vp.paragraph_format.space_before = Pt(4)
            vp.paragraph_format.space_after = Pt(4)
            vrun = vp.add_run(value)
            vrun.font.size = Pt(value_size)
            vrun.font.color.rgb = NAVY if value_bold else DARK_TEXT
            vrun.font.name = "Calibri"
            vrun.bold = value_bold

        if agency:
            submitted_to = agency
            if contracting_office:
                submitted_to += f"\n{contracting_office}"
            _add_info_row("Submitted To:", submitted_to, value_bold=True, value_size=12)
        if solicitation_number:
            _add_info_row("Solicitation No:", solicitation_number)
        if vendor_name:
            _add_info_row("Submitted By:", vendor_name, value_bold=True, value_size=12)
        if cage_code:
            _add_info_row("CAGE Code:", cage_code)
        if duns_number:
            _add_info_row("UEI / DUNS:", duns_number)
        if naics_codes:
            codes_str = ", ".join(naics_codes) if isinstance(naics_codes, list) else str(naics_codes)
            _add_info_row("NAICS Codes:", codes_str)
        _add_info_row("Submission Date:", display_date)
        if poc_name:
            poc_value = poc_name
            if poc_title_str:
                poc_value += f"\n{poc_title_str}"
            poc_details = []
            if poc_email:
                poc_details.append(poc_email)
            if poc_phone:
                poc_details.append(poc_phone)
            if poc_details:
                poc_value += f"\n{' | '.join(poc_details)}"
            _add_info_row("Point of Contact:", poc_value, value_bold=True)

        for row_idx, row in enumerate(info_table.rows):
            for cell in row.cells:
                _set_cell_borders(cell, "E5E7EB", "2")
                if row_idx % 2 == 0:
                    _set_cell_shading(cell, "F9FAFB")

        doc.add_paragraph("")

        bottom_line = doc.add_table(rows=1, cols=1)
        bottom_line.alignment = WD_TABLE_ALIGNMENT.CENTER
        blc = bottom_line.rows[0].cells[0]
        blc.text = ""
        _set_cell_shading(blc, ACCENT_HEX)
        blc.width = Inches(7)
        blp = blc.paragraphs[0]
        blp.paragraph_format.space_before = Pt(0)
        blp.paragraph_format.space_after = Pt(0)
        blr = blp.add_run(" ")
        blr.font.size = Pt(2)

        doc.add_paragraph("")

        _add_formatted_paragraph(
            doc, f"Confidential  |  {vendor_name or 'Company Name'}  |  Prepared {display_date}",
            size=9, color=RGBColor(0x6B, 0x72, 0x80), italic=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

        doc.add_page_break()

        # --- Table of Contents ---
        toc_heading = doc.add_heading("Table of Contents", level=1)
        for run in toc_heading.runs:
            run.font.color.rgb = NAVY
            run.font.name = "Calibri"

        accent_table = doc.add_table(rows=1, cols=1)
        accent_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        ac = accent_table.rows[0].cells[0]
        ac.text = ""
        ac.paragraphs[0].paragraph_format.space_before = Pt(0)
        ac.paragraphs[0].paragraph_format.space_after = Pt(0)
        _set_cell_shading(ac, ACCENT_HEX)
        ac.width = Inches(6.5)
        tr_el = accent_table.rows[0]._tr
        trPr = tr_el.get_or_add_trPr()
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="40" w:hRule="exact"/>')
        trPr.append(trHeight)

        doc.add_paragraph("")

        section_list = list(sections.items())
        total_sections = len(section_list)
        section_keys = [k for k, _ in section_list]
        volume_assignments = proposal.get("volume_assignments", {})

        if volume_assignments:
            assigned_keys = set()
            volumes = []
            for vol_label, vol_section_keys in volume_assignments.items():
                vol_items = []
                for sk in vol_section_keys:
                    if sk in sections:
                        vol_items.append((sk, sections[sk]))
                        assigned_keys.add(sk)
                if vol_items:
                    volumes.append((vol_label, vol_items))
            unassigned = [(k, v) for k, v in section_list if k not in assigned_keys]
            if unassigned:
                volumes.append(("Other Sections", unassigned))
        else:
            vol1_end = min(3, total_sections)
            vol2_end = min(vol1_end + max((total_sections - vol1_end) // 2, 3), total_sections)
            volumes = [
                ("Volume I \u2014 Administrative", section_list[0:vol1_end]),
                ("Volume II \u2014 Technical", section_list[vol1_end:vol2_end]),
                ("Volume III \u2014 Management & Compliance", section_list[vol2_end:]),
            ]

        page_num = 3
        global_idx = 0

        for vol_label, vol_items in volumes:
            if not vol_items:
                continue

            vol_para = doc.add_paragraph()
            vol_para.paragraph_format.space_before = Pt(12)
            vol_para.paragraph_format.space_after = Pt(4)
            vr = vol_para.add_run(vol_label)
            vr.font.size = Pt(12)
            vr.font.bold = True
            vr.font.color.rgb = NAVY
            vr.font.name = "Calibri"

            toc_table = doc.add_table(rows=len(vol_items), cols=3)
            toc_table.alignment = WD_TABLE_ALIGNMENT.LEFT

            for row in toc_table.rows:
                row.cells[0].width = Inches(0.4)
                row.cells[1].width = Inches(5.2)
                row.cells[2].width = Inches(0.9)

            for local_idx, (section_key, section_data) in enumerate(vol_items):
                raw_t = section_data.get("title", ""); section_title = (raw_t if raw_t and "_" not in raw_t else section_key.replace("_", " ").title())

                num_cell = toc_table.rows[local_idx].cells[0]
                num_cell.text = ""
                p = num_cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(f"{global_idx + 1}.")
                run.font.size = Pt(10)
                run.font.color.rgb = ACCENT
                run.font.bold = True
                run.font.name = "Calibri"

                title_cell = toc_table.rows[local_idx].cells[1]
                title_cell.text = ""
                p = title_cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(section_title)
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_TEXT
                run.font.name = "Calibri"

                page_cell = toc_table.rows[local_idx].cells[2]
                page_cell.text = ""
                p = page_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(str(page_num))
                run.font.size = Pt(10)
                run.font.color.rgb = GRAY_TEXT
                run.font.name = "Calibri"

                if local_idx % 2 == 0:
                    _set_cell_shading(num_cell, LIGHT_GRAY_BG)
                    _set_cell_shading(title_cell, LIGHT_GRAY_BG)
                    _set_cell_shading(page_cell, LIGHT_GRAY_BG)

                page_num += 1
                global_idx += 1

        doc.add_page_break()

        # --- Content Sections ---
        content_sections = [(k, v) for k, v in sections.items() if k != "cover_page"]
        for idx, (section_key, section_data) in enumerate(content_sections, start=1):
            raw_t = section_data.get("title", ""); section_title = (raw_t if raw_t and "_" not in raw_t else section_key.replace("_", " ").title())
            content = section_data.get("content", "")

            if idx > 1:
                doc.add_page_break()

            heading_para = doc.add_paragraph()
            heading_para.paragraph_format.space_before = Pt(0)
            heading_para.paragraph_format.space_after = Pt(2)

            section_heading = doc.add_heading(f"{idx}. {section_title}", level=1)
            for run in section_heading.runs:
                run.font.color.rgb = NAVY
                run.font.name = "Calibri"
            heading_para._element.getparent().remove(heading_para._element)

            line_table = doc.add_table(rows=1, cols=1)
            line_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell = line_table.rows[0].cells[0]
            cell.text = ""
            _set_cell_shading(cell, ACCENT_HEX)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(" ")
            run.font.size = Pt(1)

            doc.add_paragraph("")

            risk_data = _detect_risk_register(content)
            if risk_data and section_key in ('risk_mitigation', 'risk_register'):
                pre_table_lines = []
                for line in content.split('\n'):
                    if re.search(r'(?:Risk|Issue)\s*:', line, re.IGNORECASE) and '|' in line:
                        break
                    pre_table_lines.append(line)
                if pre_table_lines:
                    _render_content_lines(doc, '\n'.join(pre_table_lines))
                _add_styled_table(doc, risk_data[0], risk_data[1])
                continue

            content = _strip_leading_title(content, section_title)
            _render_content_lines(doc, content, section_key)

        doc.add_paragraph("")
        _add_formatted_paragraph(
            doc,
            f"Generated by GovProposal AI on {datetime.now().strftime('%B %d, %Y')}",
            size=8, color=GRAY_TEXT, italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as exc:
        logger.error("Failed to generate DOCX: %s", exc)
        raise RuntimeError(f"DOCX generation failed: {exc}")


def _render_content_lines(doc, content: str, section_key: str = ""):
    """Render content with smart detection of tables, lists, and structured data."""
    parsed = _html_to_paragraphs(content)

    i = 0
    while i < len(parsed):
        para = parsed[i]
        ptype = para['type']
        ptext = para['text']

        if ptype == 'spacer':
            doc.add_paragraph("")
            i += 1
            continue

        if ptype == 'html_table':
            td = para.get('table_data', {})
            headers = td.get('headers', [])
            data_rows = td.get('rows', [])
            if headers or data_rows:
                _add_html_table_to_docx(doc, headers, data_rows)
                doc.add_paragraph("")
            i += 1
            continue

        if '|' in ptext and ptext.count('|') >= 2:
            table_lines = [ptext]
            j = i + 1
            while j < len(parsed) and ('|' in parsed[j]['text'] and parsed[j]['text'].count('|') >= 2 or re.match(r'^[-|:]+$', parsed[j]['text'].strip())):
                table_lines.append(parsed[j]['text'])
                j += 1
            if len(table_lines) >= 2:
                table_data = _detect_table_content(table_lines)
                if table_data:
                    _add_styled_table(doc, table_data[0], table_data[1])
                    doc.add_paragraph("")
                    i = j
                    continue

        if ptype in ('h1', 'h2'):
            heading = doc.add_heading(_strip_markdown_bold(ptext), level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
                run.font.name = "Calibri"
        elif ptype == 'h3':
            heading = doc.add_heading(_strip_markdown_bold(ptext), level=3)
            for run in heading.runs:
                run.font.color.rgb = DARK_TEXT
                run.font.name = "Calibri"
        elif ptype == 'bullet':
            clean = _strip_markdown_bold(ptext)
            if clean:
                bp = doc.add_paragraph(style="List Bullet")
                parts = re.split(r'(\*\*.*?\*\*)', ptext)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = bp.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = bp.add_run(part)
                    run.font.size = Pt(11)
                    run.font.color.rgb = DARK_TEXT
                    run.font.name = "Calibri"
        elif ptype == 'numbered':
            clean = _strip_markdown_bold(ptext)
            if clean:
                np = doc.add_paragraph(style="List Number")
                run = np.add_run(clean)
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_TEXT
                run.font.name = "Calibri"
        else:
            clean = ptext.strip()
            if clean:
                _add_formatted_paragraph(doc, clean, size=11, color=DARK_TEXT)

        i += 1


# ============================================================
# PDF Export
# ============================================================

def generate_pdf(proposal: Dict) -> io.BytesIO:
    """Generate a professional PDF document from proposal sections."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.85 * inch,
        )

        # ── Dark / Black theme colours ──────────────────────────────────
        BG_DARK      = "#0F1117"   # page background (simulated via section blocks)
        CARD_DARK    = "#1A1D27"   # section card bg
        ACCENT_GOLD  = "#F5A623"   # gold accent
        HEADING_WHITE= "#FFFFFF"
        BODY_LIGHT   = "#D1D5DB"
        SUBHEAD_GOLD = "#F5A623"
        BORDER_DARK  = "#2D3148"

        # Override template colours with dark theme
        heading_hex  = "1A1D27"
        accent_hex   = "F5A623"

        styles = getSampleStyleSheet()

        # Helper: no-hyphen body style
        def _body(name, **kw):
            defaults = dict(
                parent=styles["Normal"],
                fontSize=10.5, leading=16, spaceAfter=7,
                fontName="Helvetica",
                textColor=colors.HexColor(BODY_LIGHT),
                wordWrap='LTR',
                hyphenationLang='',          # disable hyphenation
                embeddedHyphenation=0,
                splitLongWords=0,
                alignment=4,                 # JUSTIFY
            )
            defaults.update(kw)
            return ParagraphStyle(name, **defaults)

        styles.add(ParagraphStyle(
            name="ProposalTitle", parent=styles["Title"],
            fontSize=26, textColor=colors.HexColor(HEADING_WHITE),
            spaceAfter=12, alignment=1, fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="ProposalSubtitle", parent=styles["Normal"],
            fontSize=14, textColor=colors.HexColor(ACCENT_GOLD),
            spaceAfter=8, alignment=1, fontName="Helvetica",
        ))
        styles.add(ParagraphStyle(
            name="SectionHeading", parent=styles["Heading1"],
            fontSize=16, textColor=colors.HexColor(HEADING_WHITE),
            spaceBefore=20, spaceAfter=10, fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="SubHeading", parent=styles["Heading2"],
            fontSize=12, textColor=colors.HexColor(ACCENT_GOLD),
            spaceBefore=10, spaceAfter=5, fontName="Helvetica-Bold",
        ))
        styles.add(_body("ProposalBody"))
        styles.add(_body(
            "ProposalBullet",
            leftIndent=18, bulletIndent=8, spaceAfter=4,
        ))
        styles.add(ParagraphStyle(
            name="Confidential", parent=styles["Normal"],
            fontSize=9, textColor=colors.HexColor("#999999"),
            alignment=1, fontName="Helvetica-Oblique",
        ))
        styles.add(ParagraphStyle(
            name="TableHeader", parent=styles["Normal"],
            fontSize=9, textColor=colors.white,
            fontName="Helvetica-Bold", alignment=1,
        ))
        styles.add(ParagraphStyle(
            name="TableCell", parent=styles["Normal"],
            fontSize=9, textColor=colors.HexColor(BODY_LIGHT),
            fontName="Helvetica", leading=13,
            wordWrap='LTR', splitLongWords=0,
        ))


        story = []
        proposal_title = proposal.get("proposal_title", "Government Proposal")
        vendor_name = proposal.get("vendor_name", "")
        sections = proposal.get("sections", {})
        template = proposal.get("template") or {}
        accent_hex = _normalize_hex_color(template.get("accent", ""), "#10B981")
        heading_hex = _normalize_hex_color(template.get("headingColor", ""), "#1B2A4A")
        styles["ProposalTitle"].textColor = colors.HexColor("#0F1117")
        styles["ProposalSubtitle"].textColor = colors.HexColor("#F5A623")
        styles["SectionHeading"].textColor = colors.HexColor("#0F1117")
        styles["SubHeading"].textColor = colors.HexColor("#F5A623")

        meta = proposal.get("metadata") or {}
        proposal_type = meta.get("proposal_type", "Government Contract Proposal")
        agency = meta.get("agency", "")
        contracting_office = meta.get("contracting_office", "")
        solicitation_number = meta.get("solicitation_number", "")
        submission_date = meta.get("submission_date", "")
        cage_code = meta.get("cage_code", "")
        duns_number = meta.get("duns_number", "")
        naics_codes = meta.get("naics_codes", [])
        poc_name = meta.get("poc_name", "")
        poc_title_str = meta.get("poc_title", "")
        poc_email = meta.get("poc_email", "")
        poc_phone = meta.get("poc_phone", "")
        display_date = submission_date if submission_date else datetime.now().strftime('%B %d, %Y')

        label_style = ParagraphStyle(
            "CoverLabel", parent=styles["Normal"],
            fontSize=10, textColor=colors.HexColor("#999999"),
            alignment=1, fontName="Helvetica",
        )

        # Dark background block at top
        story.append(RLTable([[""]], colWidths=[PDF_CONTENT_WIDTH], rowHeights=[0.12 * inch]))

        # --- Cover Page ---
        banner_data = [[RLParagraph(
            f"<b>{_escape_xml(proposal_title.upper())}</b>",
            ParagraphStyle("BannerTitle", parent=styles["Normal"],
                           fontSize=22, textColor=colors.white, alignment=1,
                           fontName="Helvetica-Bold", leading=28)
        )]]
        banner = RLTable(banner_data, colWidths=[PDF_CONTENT_WIDTH], rowHeights=[1.0 * inch])
        banner.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0F1117")),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 20),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 20),
            ("LEFTPADDING", (0, 0), (-1, -1), 20),
        ]))
        story.append(banner)

        accent_data = [[RLParagraph(
            f"<b>{_escape_xml(proposal_type)}</b>",
            ParagraphStyle("AccentBar", parent=styles["Normal"],
                           fontSize=12, textColor=colors.HexColor("#0F1117"), alignment=1,
                           fontName="Helvetica-Bold")
        )]]
        accent_bar = RLTable(accent_data, colWidths=[PDF_CONTENT_WIDTH], rowHeights=[0.4 * inch])
        accent_bar.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5A623")),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(accent_bar)

        story.append(Spacer(1, 0.3 * inch))

        # Company logo — constrained to page width
        company_logo = proposal.get("company_logo", "")
        if company_logo and company_logo.startswith("data:image"):
            try:
                from reportlab.platypus import Image as RLImage
                header, b64data = company_logo.split(",", 1)
                logo_bytes = base64.b64decode(b64data)
                logo_stream = io.BytesIO(logo_bytes)
                from PIL import Image as PILImage
                _pil_img = PILImage.open(io.BytesIO(logo_bytes))
                _iw, _ih = _pil_img.size
                _max_w, _max_h = 2.0, 1.0
                _scale = min(_max_w / (_iw / 96.0), _max_h / (_ih / 96.0), 1.0)
                _fit_w = (_iw / 96.0) * _scale
                _fit_h = (_ih / 96.0) * _scale
                logo_img = RLImage(logo_stream, width=_fit_w * inch, height=_fit_h * inch)
                logo_table = RLTable([[logo_img]], colWidths=[PDF_CONTENT_WIDTH])
                logo_table.setStyle(TableStyle([
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]))
                story.append(logo_table)
            except Exception as logo_err:
                logger.warning("Failed to embed logo in PDF: %s", logo_err)
                story.append(RLParagraph("[Company Logo]", label_style))
        else:
            story.append(RLParagraph("[Upload logo in Business Profile]", label_style))

        story.append(Spacer(1, 0.3 * inch))

        # Info table
        info_rows = []
        if agency:
            agency_text = _escape_xml(agency)
            if contracting_office:
                agency_text += f"<br/>{_escape_xml(contracting_office)}"
            info_rows.append(["Submitted To:", agency_text, True])
        if solicitation_number:
            info_rows.append(["Solicitation No:", _escape_xml(solicitation_number), False])
        if vendor_name:
            info_rows.append(["Submitted By:", _escape_xml(vendor_name), True])
        if cage_code:
            info_rows.append(["CAGE Code:", _escape_xml(cage_code), False])
        if duns_number:
            info_rows.append(["UEI / DUNS:", _escape_xml(duns_number), False])
        if naics_codes:
            codes_str = ", ".join(naics_codes) if isinstance(naics_codes, list) else str(naics_codes)
            info_rows.append(["NAICS Codes:", _escape_xml(codes_str), False])
        info_rows.append(["Submission Date:", _escape_xml(display_date), False])
        if poc_name:
            poc_val = f"<b>{_escape_xml(poc_name)}</b>"
            if poc_title_str:
                poc_val += f"<br/>{_escape_xml(poc_title_str)}"
            poc_parts = []
            if poc_email:
                poc_parts.append(_escape_xml(poc_email))
            if poc_phone:
                poc_parts.append(_escape_xml(poc_phone))
            if poc_parts:
                poc_val += f"<br/>{' | '.join(poc_parts)}"
            info_rows.append(["Point of Contact:", poc_val, False])

        if info_rows:
            table_data = []
            for lbl, val, val_bold in info_rows:
                lbl_para = RLParagraph(f"<b>{_escape_xml(lbl)}</b>", ParagraphStyle(
                    "InfoLabel", parent=styles["Normal"], fontSize=9.5,
                    textColor=colors.HexColor("#F5A623"), alignment=2, fontName="Helvetica-Bold"))
                val_text = f"<b>{val}</b>" if val_bold else val
                val_para = RLParagraph(val_text, ParagraphStyle(
                    "InfoValue", parent=styles["Normal"], fontSize=10.5,
                    textColor=colors.white if val_bold else colors.HexColor("#2D3148"),
                    fontName="Helvetica-Bold" if val_bold else "Helvetica", leading=14,
                    wordWrap='LTR', splitLongWords=0))
                table_data.append([lbl_para, val_para])

            info_tbl = RLTable(table_data, colWidths=[1.8 * inch, PDF_CONTENT_WIDTH - 1.8 * inch])
            tbl_style = [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (0, -1), 12),
                ("RIGHTPADDING", (0, 0), (0, -1), 8),
                ("LEFTPADDING", (1, 0), (1, -1), 8),
                ("LINEBELOW", (0, 0), (-1, -2), 0.5, colors.HexColor("#2D3148")),
            ]
            for i in range(0, len(table_data), 2):
                tbl_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#1A1D27")))
            info_tbl.setStyle(TableStyle(tbl_style))
            story.append(info_tbl)

        story.append(Spacer(1, 0.4 * inch))

        bottom = RLTable([[""]], colWidths=[PDF_CONTENT_WIDTH], rowHeights=[4])
        bottom.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5A623")),
        ]))
        story.append(bottom)

        story.append(Spacer(1, 0.2 * inch))
        story.append(RLParagraph(
            f"Confidential  |  {_escape_xml(vendor_name or 'Company Name')}  |  Prepared {_escape_xml(display_date)}",
            ParagraphStyle("CoverFooterNote", parent=styles["Normal"],
                           fontSize=9, textColor=colors.HexColor("#999999"),
                           alignment=1, fontName="Helvetica")))

        # NOTE: floating_images intentionally excluded from PDF to prevent layout overflow

        story.append(PageBreak())

        # --- Table of Contents ---
        story.append(RLParagraph("Table of Contents", styles["SectionHeading"]))

        toc_line = RLTable([[""]], colWidths=[PDF_CONTENT_WIDTH], rowHeights=[2])
        toc_line.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5A623")),
        ]))
        story.append(toc_line)
        story.append(Spacer(1, 0.15 * inch))

        section_list = list(sections.items())
        total_sections = len(section_list)
        volume_assignments = proposal.get("volume_assignments", {})

        if volume_assignments:
            assigned_keys = set()
            pdf_vol_list = []
            for vol_label, vol_section_keys in volume_assignments.items():
                vol_items = []
                for sk in vol_section_keys:
                    if sk in sections:
                        vol_items.append((sk, sections[sk]))
                        assigned_keys.add(sk)
                if vol_items:
                    pdf_vol_list.append((vol_label, vol_items))
            unassigned = [(k, v) for k, v in section_list if k not in assigned_keys]
            if unassigned:
                pdf_vol_list.append(("Other Sections", unassigned))
        else:
            vol1_end = min(3, total_sections)
            vol2_end = min(vol1_end + max((total_sections - vol1_end) // 2, 3), total_sections)
            pdf_vol_list = [
                ("Volume I \u2014 Administrative", section_list[0:vol1_end]),
                ("Volume II \u2014 Technical", section_list[vol1_end:vol2_end]),
                ("Volume III \u2014 Management & Compliance", section_list[vol2_end:]),
            ]

        vol_heading_style = ParagraphStyle(
            "VolHeading", parent=styles["Normal"],
            fontSize=11, textColor=colors.HexColor("#0F1117"),
            fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4,
        )
        page_style = ParagraphStyle(
            "PageNum", parent=styles["TableCell"],
            alignment=2,
            textColor=colors.HexColor("#999999"),
        )

        page_num = 3
        global_idx = 0
        for vol_label, vol_items in pdf_vol_list:
            if not vol_items:
                continue

            story.append(RLParagraph(_escape_xml(vol_label), vol_heading_style))

            toc_data = []
            for local_idx, (section_key, section_data) in enumerate(vol_items):
                raw_t = section_data.get("title", ""); section_title = (raw_t if raw_t and "_" not in raw_t else section_key.replace("_", " ").title())
                toc_data.append([
                    RLParagraph(f"<b>{global_idx + 1}.</b>", styles["TableCell"]),
                    RLParagraph(_escape_xml(section_title), styles["TableCell"]),
                    RLParagraph(str(page_num), page_style),
                ])
                page_num += 1
                global_idx += 1

            if toc_data:
                toc_table = RLTable(toc_data, colWidths=[0.35 * inch, PDF_CONTENT_WIDTH - 1.15 * inch, 0.8 * inch])
                toc_style = [
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
                for i in range(0, len(toc_data), 2):
                    toc_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#1A1D27")))
                toc_table.setStyle(TableStyle(toc_style))
                story.append(toc_table)

        story.append(PageBreak())

        # --- Content Sections ---
        pdf_content_sections = [(k, v) for k, v in sections.items() if k != "cover_page"]
        for idx, (section_key, section_data) in enumerate(pdf_content_sections, start=1):
            raw_t = section_data.get("title", ""); section_title = (raw_t if raw_t and "_" not in raw_t else section_key.replace("_", " ").title())
            content = section_data.get("content", "")

            if idx > 1:
                story.append(PageBreak())

            story.append(RLParagraph(
                f"{idx}. {_escape_xml(section_title)}",
                styles["SectionHeading"],
            ))

            line = RLTable([[""]], colWidths=[PDF_CONTENT_WIDTH], rowHeights=[2])
            line.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5A623")),
            ]))
            story.append(line)
            story.append(Spacer(1, 0.15 * inch))

            content = _strip_leading_title(content, section_title)

            # Render section images (constrained to page width)
            section_images = section_data.get("images") or []
            for img in section_images[:3]:
                img_url = img.get("url", "") if isinstance(img, dict) else img
                if img_url:
                    _add_pdf_image(story, img_url, styles)

            parsed = _html_to_paragraphs(content)
            _render_pdf_content(story, parsed, styles, section_key)
            story.append(Spacer(1, 0.3 * inch))

        # Footer
        story.append(Spacer(1, 0.5 * inch))
        footer_style = ParagraphStyle(
            "Footer", parent=styles["Normal"],
            fontSize=8, textColor=colors.HexColor("#999999"),
            alignment=1, fontName="Helvetica-Oblique",
        )
        story.append(RLParagraph(
            f"Generated by GovProposal AI on {datetime.now().strftime('%B %d, %Y')}",
            footer_style
        ))

        # Footer callback — dark background + gold accents on every page
        company_logo = proposal.get("company_logo", "")
        _pdf_logo_bytes = None
        if company_logo and company_logo.startswith("data:image"):
            try:
                _, b64data = company_logo.split(",", 1)
                _pdf_logo_bytes = base64.b64decode(b64data)
            except Exception:
                pass

        def _draw_footer(canvas, doc_obj):
            canvas.saveState()
            page_w, page_h = letter

            # Dark background for entire page
            canvas.setFillColor(colors.HexColor("#0F1117"))
            canvas.rect(0, 0, page_w, page_h, fill=1, stroke=0)

            # Gold top accent bar
            canvas.setFillColor(colors.HexColor("#F5A623"))
            canvas.rect(0, page_h - 4, page_w, 4, fill=1, stroke=0)

            # Footer bar
            footer_y = 0.35 * inch
            canvas.setFillColor(colors.HexColor("#1A1D27"))
            canvas.rect(0, 0, page_w, 0.6 * inch, fill=1, stroke=0)

            if _pdf_logo_bytes:
                try:
                    from PIL import Image as PILImage
                    from reportlab.lib.utils import ImageReader
                    _pil = PILImage.open(io.BytesIO(_pdf_logo_bytes))
                    _iw, _ih = _pil.size
                    _max_w, _max_h = 0.8, 0.28
                    _scale = min(_max_w / (_iw / 96.0), _max_h / (_ih / 96.0), 1.0)
                    _fw = (_iw / 96.0) * _scale * inch
                    _fh = (_ih / 96.0) * _scale * inch
                    img_reader = ImageReader(io.BytesIO(_pdf_logo_bytes))
                    canvas.drawImage(img_reader, 0.75 * inch, footer_y - _fh / 2,
                                     width=_fw, height=_fh, mask='auto')
                except Exception:
                    canvas.setFont("Helvetica", 7)
                    canvas.setFillColor(colors.HexColor("#999999"))
                    canvas.drawString(0.75 * inch, footer_y, vendor_name or "")
            else:
                canvas.setFont("Helvetica", 7)
                canvas.setFillColor(colors.HexColor("#999999"))
                canvas.drawString(0.75 * inch, footer_y, vendor_name or "")

            canvas.setFont("Helvetica-Bold", 8)
            canvas.setFillColor(colors.HexColor("#F5A623"))
            canvas.drawCentredString(page_w / 2, footer_y, f"Page {canvas.getPageNumber()}")

            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(colors.HexColor("#999999"))
            canvas.drawRightString(page_w - 0.75 * inch, footer_y,
                                   f"Confidential  |  {vendor_name or ''}")

            canvas.restoreState()

        doc.build(story, onFirstPage=_draw_footer, onLaterPages=_draw_footer)
        buffer.seek(0)
        return buffer

    except Exception as exc:
        logger.error("Failed to generate PDF: %s", exc)
        raise RuntimeError(f"PDF generation failed: {exc}")


def _render_pdf_content(story, parsed, styles, section_key=""):
    """Render parsed content for PDF with smart table detection."""
    i = 0
    while i < len(parsed):
        para = parsed[i]
        ptype = para['type']
        ptext = para['text']

        if ptype == 'spacer':
            story.append(Spacer(1, 4))
            i += 1
            continue

        if ptype == 'html_table':
            td = para.get('table_data', {})
            headers = td.get('headers', [])
            data_rows = td.get('rows', [])
            if headers or data_rows:
                _add_html_table_to_pdf(story, headers, data_rows, styles)
                story.append(Spacer(1, 8))
            i += 1
            continue

        if ptype == 'image':
            _add_pdf_image(story, para.get('src', ''), styles)
            i += 1
            continue

        if '|' in ptext and ptext.count('|') >= 2:
            table_lines = [ptext]
            j = i + 1
            while j < len(parsed) and ('|' in parsed[j]['text'] and parsed[j]['text'].count('|') >= 2 or re.match(r'^[-|:]+$', parsed[j]['text'].strip())):
                table_lines.append(parsed[j]['text'])
                j += 1
            if len(table_lines) >= 2:
                table_data = _detect_table_content(table_lines)
                if table_data:
                    _add_pdf_table(story, table_data[0], table_data[1], styles)
                    story.append(Spacer(1, 8))
                    i = j
                    continue

        if ptype in ('h1', 'h2'):
            story.append(RLParagraph(_escape_xml(ptext), styles["SubHeading"]))
        elif ptype == 'h3':
            story.append(RLParagraph(f"<b>{_escape_xml(ptext)}</b>", styles["ProposalBody"]))
        elif ptype == 'bullet':
            bullet_text = _markdown_bold_to_reportlab(ptext)
            story.append(RLParagraph(f"\u2022  {bullet_text}", styles["ProposalBullet"]))
        else:
            formatted = _markdown_bold_to_reportlab(ptext)
            if formatted.strip():
                story.append(RLParagraph(formatted, styles["ProposalBody"]))

        i += 1


def _add_pdf_table(story, headers, rows, styles):
    """Add a styled table to PDF story, constrained to page width."""
    header_row = [RLParagraph(f"<b>{_escape_xml(h)}</b>", styles["TableHeader"]) for h in headers]
    data = [header_row]
    for row in rows:
        data.append([RLParagraph(_escape_xml(str(c)), styles["TableCell"]) for c in row])

    num_cols = len(headers)
    col_width = PDF_CONTENT_WIDTH / num_cols

    table = RLTable(data, colWidths=[col_width] * num_cols)
    table_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F1117")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#2D3148")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("WORDWRAP", (0, 0), (-1, -1), True),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0:
            table_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#1A1D27")))

    table.setStyle(TableStyle(table_style))
    story.append(table)


def _load_image_bytes(src: str) -> Optional[bytes]:
    """Load an image from a data URL, uploaded asset path, local file, or public URL."""
    if not src:
        return None

    if src.startswith("data:image"):
        try:
            _, b64data = src.split(",", 1)
            return base64.b64decode(b64data)
        except Exception as exc:
            logger.warning("Failed to decode image data URL: %s", exc)
            return None

    parsed = urllib.parse.urlparse(src)
    asset_path = parsed.path if parsed.scheme else src
    if asset_path.startswith("/api/uploads/"):
        parts = [p for p in asset_path.split("/") if p]
        if len(parts) >= 4:
            file_path = UPLOADS_DIR / parts[-2] / parts[-1]
            try:
                if file_path.exists():
                    return file_path.read_bytes()
            except Exception as exc:
                logger.warning("Failed to read uploaded image %s: %s", file_path, exc)

    if parsed.scheme in {"http", "https"}:
        try:
            with urllib.request.urlopen(src, timeout=8) as response:
                return response.read(8 * 1024 * 1024)
        except Exception as exc:
            logger.warning("Failed to fetch image %s: %s", src, exc)
            return None

    return None


def _add_pdf_image(story, src: str, styles):
    """Add an image to the PDF story, strictly constrained to page width."""
    image_bytes = _load_image_bytes(src)
    if not image_bytes:
        return

    try:
        from PIL import Image as PILImage
        from reportlab.platypus import Image as RLImage

        pil_image = PILImage.open(io.BytesIO(image_bytes))
        width_px, height_px = pil_image.size
        if width_px <= 0 or height_px <= 0:
            return

        # Strictly constrain to content width — never exceed page margins
        max_w = PDF_CONTENT_WIDTH
        max_h = 3.0 * inch
        width_pt = width_px / 96.0 * inch
        height_pt = height_px / 96.0 * inch
        scale = min(max_w / width_pt, max_h / height_pt, 1.0)

        img = RLImage(io.BytesIO(image_bytes), width=width_pt * scale, height=height_pt * scale)
        img_table = RLTable([[img]], colWidths=[PDF_CONTENT_WIDTH])
        img_table.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(img_table)
        story.append(Spacer(1, 8))
    except Exception as exc:
        logger.warning("Failed to render image in PDF: %s", exc)


def _add_html_table_to_pdf(story, headers: List[str], rows: List[Dict], styles):
    """Render an HTML table as a styled PDF table, constrained to page width."""
    if not headers and not rows:
        return

    num_cols = len(headers) if headers else (len(rows[0]['cells']) if rows else 0)
    if num_cols == 0:
        return

    data = []
    summary_row_indices = []

    if headers:
        header_row = [RLParagraph(f"<b>{_escape_xml(h)}</b>", styles["TableHeader"]) for h in headers]
        data.append(header_row)

    for row_idx, row_data in enumerate(rows):
        cells_data = row_data.get('cells', [])
        is_summary = row_data.get('is_summary', False)
        actual_row_idx = len(data)

        cell_style = styles["TableCell"]
        row_cells = []
        for col_idx in range(num_cols):
            cell_text = cells_data[col_idx] if col_idx < len(cells_data) else ""
            if is_summary:
                row_cells.append(RLParagraph(f"<b>{_escape_xml(cell_text)}</b>", cell_style))
            else:
                row_cells.append(RLParagraph(_escape_xml(cell_text), cell_style))
        data.append(row_cells)

        if is_summary:
            summary_row_indices.append(actual_row_idx)

    col_width = PDF_CONTENT_WIDTH / num_cols

    table = RLTable(data, colWidths=[col_width] * num_cols)
    table_style = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#2D3148")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("WORDWRAP", (0, 0), (-1, -1), True),
    ]

    if headers:
        table_style.extend([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F1117")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ])

    for i in range(1 if headers else 0, len(data)):
        if i in summary_row_indices:
            table_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#1A1D27")))
        elif i % 2 == 0:
            table_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#1A1D27")))

    table.setStyle(TableStyle(table_style))
    story.append(table)


# ============================================================
# Utility functions
# ============================================================

def _escape_xml(text: str) -> str:
    """Escape special XML characters for ReportLab Paragraph."""
    text = str(text)
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    return text


def _strip_markdown_bold(text: str) -> str:
    """Remove markdown bold markers (**text**) for DOCX output."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _markdown_bold_to_reportlab(text: str) -> str:
    """Convert markdown bold (**text**) to ReportLab <b>text</b> tags."""
    text = _escape_xml(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return text


def _parse_html_tables(html_content: str) -> List[Dict]:
    """Extract HTML tables from content."""
    tables = []
    table_pattern = re.compile(r'<table[^>]*>(.*?)</table>', re.DOTALL | re.IGNORECASE)

    for table_match in table_pattern.finditer(html_content):
        table_html = table_match.group(1)
        headers = []
        rows = []

        th_pattern = re.compile(r'<th[^>]*>(.*?)</th>', re.DOTALL | re.IGNORECASE)
        header_row_match = re.search(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
        if header_row_match:
            ths = th_pattern.findall(header_row_match.group(1))
            if ths:
                headers = [re.sub(r'<[^>]+>', '', h).strip() for h in ths]

        tr_pattern = re.compile(r'<tr[^>]*>(.*?)</tr>', re.DOTALL | re.IGNORECASE)
        td_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.DOTALL | re.IGNORECASE)

        for i, tr_match in enumerate(tr_pattern.finditer(table_html)):
            row_html = tr_match.group(1)
            if i == 0 and th_pattern.search(row_html):
                continue
            cells = td_pattern.findall(row_html)
            if cells:
                cleaned = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
                is_summary = bool(re.search(r'colspan|font-weight:\s*bold', tr_match.group(0), re.IGNORECASE))
                rows.append({'cells': cleaned, 'is_summary': is_summary})

        if headers or rows:
            tables.append({'headers': headers, 'rows': rows})

    return tables


def _html_to_paragraphs(html_content: str) -> list:
    """Convert HTML content into structured paragraph dicts."""
    if not html_content:
        return []

    paragraphs = []

    tables = _parse_html_tables(html_content)
    images = []
    text = html_content
    table_idx = 0
    table_pattern = re.compile(r'<table[^>]*>.*?</table>', re.DOTALL | re.IGNORECASE)
    for match in table_pattern.finditer(html_content):
        text = text.replace(match.group(0), f'\n__HTML_TABLE_{table_idx}__\n', 1)
        table_idx += 1

    image_pattern = re.compile(r'<img\b[^>]*\bsrc=["\']([^"\']+)["\'][^>]*>', re.IGNORECASE)
    image_idx = 0
    for match in list(image_pattern.finditer(text)):
        src = html_unescape(match.group(1).strip())
        images.append({'src': src})
        text = text.replace(match.group(0), f'\n__HTML_IMAGE_{image_idx}__\n', 1)
        image_idx += 1

    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<(?:strong|b)>(.*?)</(?:strong|b)>', r'**\1**', text, flags=re.DOTALL)
    text = re.sub(r'<(?:em|i)>(.*?)</(?:em|i)>', r'_\1_', text, flags=re.DOTALL)
    text = re.sub(r'<li>(.*?)</li>', r'\n- \1\n', text, flags=re.DOTALL)

    for level in range(1, 4):
        text = re.sub(
            rf'<h{level}[^>]*>(.*?)</h{level}>',
            lambda m, l=level: f'\n{"#" * l} {m.group(1).strip()}\n',
            text, flags=re.DOTALL,
        )

    text = re.sub(r'<[^>]+>', '', text)
    text = html_unescape(text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            paragraphs.append({'type': 'spacer', 'text': ''})
        elif re.match(r'^__HTML_TABLE_(\d+)__$', line):
            idx = int(re.match(r'^__HTML_TABLE_(\d+)__$', line).group(1))
            if idx < len(tables):
                paragraphs.append({'type': 'html_table', 'text': '', 'table_data': tables[idx]})
        elif re.match(r'^__HTML_IMAGE_(\d+)__$', line):
            idx = int(re.match(r'^__HTML_IMAGE_(\d+)__$', line).group(1))
            if idx < len(images):
                paragraphs.append({'type': 'image', 'text': '', 'src': images[idx]['src']})
        elif line.startswith('### '):
            paragraphs.append({'type': 'h3', 'text': line[4:].strip()})
        elif line.startswith('## '):
            paragraphs.append({'type': 'h2', 'text': line[3:].strip()})
        elif line.startswith('# '):
            paragraphs.append({'type': 'h1', 'text': line[2:].strip()})
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('\u2022 '):
            prefix_len = 2
            paragraphs.append({'type': 'bullet', 'text': line[prefix_len:].strip()})
        elif re.match(r'^\d+\.\s', line):
            paragraphs.append({'type': 'numbered', 'text': re.sub(r'^\d+\.\s*', '', line).strip()})
        else:
            paragraphs.append({'type': 'text', 'text': line})

    return paragraphs